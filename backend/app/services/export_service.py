"""
Сервис для экспорта данных персонажей в различных форматах.
"""

import io
import os
import time
from typing import Dict, Any, Optional, BinaryIO
from datetime import datetime
from pathlib import Path

# PDF генерация
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# DOCX генерация
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

# HTML to PDF (опциональное)
try:
    import weasyprint
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False
    weasyprint = None

from jinja2 import Template, Environment, FileSystemLoader

from app.database.models.character import Character
from app.database.models.checklist import ChecklistResponse
from app.schemas.checklist import ChecklistWithResponses
from app.config.settings import settings
from app.utils.error_handlers import ExportError, ErrorHandler, ErrorCode
from app.utils.logging_config import LoggingConfig


class ExportService:
    """Сервис для экспорта данных персонажей в PDF и DOCX форматы."""
    
    def __init__(self):
        """Инициализация сервиса экспорта."""
        self.template_dir = Path(__file__).parent.parent / "templates" / "export"
        self.template_dir.mkdir(parents=True, exist_ok=True)
        
        # Настройка Jinja2 environment
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.template_dir)),
            autoescape=True
        )
        
        # Проверяем доступность WeasyPrint
        if not WEASYPRINT_AVAILABLE:
            LoggingConfig.get_export_logger().warning(
                "WeasyPrint not available - HTML to PDF conversion disabled"
            )
        
        # Попытка зарегистрировать шрифт поддерживающий кириллицу для ReportLab
        try:
            font_path = self._find_cyrillic_font()
            if font_path:
                pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))
                self.default_font = 'DejaVuSans'
            else:
                self.default_font = 'Helvetica'
        except Exception:
            self.default_font = 'Helvetica'
    
    def _find_cyrillic_font(self) -> Optional[str]:
        """Поиск шрифта с поддержкой кириллицы."""
        # Общие пути к шрифтам в разных ОС
        font_paths = [
            '/System/Library/Fonts/Helvetica.ttc',  # macOS
            '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',  # Ubuntu
            'C:\\Windows\\Fonts\\arial.ttf',  # Windows
            '/usr/share/fonts/TTF/DejaVuSans.ttf',  # Arch Linux
        ]
        
        for path in font_paths:
            if os.path.exists(path):
                return path
        return None
    
    async def export_character_pdf(
        self,
        character: Character,
        checklists: list[ChecklistWithResponses],
        format_type: str = "detailed",
        user_id: Optional[int] = None,
        use_weasyprint: bool = True,
        theme: str = "default",
        custom_fonts: Optional[Dict[str, str]] = None
    ) -> bytes:
        """
        Экспорт персонажа в PDF формат.
        
        Args:
            character: Модель персонажа
            checklists: Список чеклистов с ответами
            format_type: Тип отчета ('detailed', 'summary', 'compact')
            user_id: ID пользователя для логирования
            use_weasyprint: Использовать WeasyPrint вместо ReportLab
            theme: Тема оформления ('default', 'professional', 'creative', 'minimal')
            custom_fonts: Словарь кастомных шрифтов {имя: путь_к_файлу}
            
        Returns:
            PDF файл в виде байтов
        """
        start_time = time.time()
        
        try:
            # Используем WeasyPrint если доступен и запрошен
            if use_weasyprint and WEASYPRINT_AVAILABLE:
                return await self._export_pdf_with_weasyprint(
                    character, checklists, format_type, user_id, start_time, theme, custom_fonts
                )
            else:
                return await self._export_pdf_with_reportlab(
                    character, checklists, format_type, user_id, start_time
                )
                
        except Exception as e:
            duration_ms = (time.time() - start_time) * 1000
            LoggingConfig.log_export_operation(
                operation="pdf_export",
                character_id=character.id,
                format_type=format_type,
                export_type="pdf",
                user_id=user_id,
                duration_ms=duration_ms,
                success=False,
                error=str(e)
            )
            
            if "font" in str(e).lower():
                raise ExportError(
                    "Ошибка настройки шрифтов PDF",
                    details=f"Проблема с шрифтами: {str(e)}"
                )
            elif "weasyprint" in str(e).lower():
                raise ExportError(
                    "Ошибка WeasyPrint генерации PDF",
                    details=f"WeasyPrint ошибка: {str(e)}"
                )
            elif "reportlab" in str(e).lower():
                raise ExportError(
                    "Ошибка генерации PDF документа",
                    details=f"ReportLab ошибка: {str(e)}"
                )
            else:
                raise ExportError(
                    "Неизвестная ошибка при создании PDF",
                    details=str(e)
                )
    
    async def _export_pdf_with_weasyprint(
        self,
        character: Character,
        checklists: list[ChecklistWithResponses],
        format_type: str,
        user_id: Optional[int],
        start_time: float,
        theme: str = "default",
        custom_fonts: Optional[Dict[str, str]] = None
    ) -> bytes:
        """Экспорт PDF с использованием WeasyPrint."""
        try:
            # Преобразуем format_type в строку если это enum
            format_type_str = format_type
            if hasattr(format_type, 'value'):
                format_type_str = format_type.value.lower()
            elif hasattr(format_type, 'name'):
                format_type_str = format_type.name.lower()
            else:
                format_type_str = str(format_type).lower()
            
            # Определяем шаблон в зависимости от типа
            template_name = f"character_{format_type_str}_weasy.html"
            
            # Проверяем существование шаблона
            template_path = self.template_dir / template_name
            if not template_path.exists():
                # Fallback на обычный шаблон
                template_name = f"character_{format_type_str}.html"
            
            template = self.jinja_env.get_template(template_name)
            
            # Определяем пол персонажа
            character_gender = self._detect_character_gender(character)
            
            # Загружаем тему если указана
            theme_css = self._load_theme_css(theme)
            
            # Подготавливаем контекст для шаблона
            context = {
                'character': character,
                'checklists': checklists,
                'export_date': datetime.now().strftime("%d.%m.%Y %H:%M"),
                'format_type': format_type,
                'character_gender': character_gender,
                'theme_css': theme_css,
                'custom_fonts': custom_fonts or {},
                'get_answer_text': lambda response, gender: self._get_answer_text(response, gender)
            }
            
            # Рендерим HTML
            html_content = template.render(**context)
            
            # Генерируем PDF с помощью WeasyPrint
            html_doc = weasyprint.HTML(string=html_content)
            
            # Настройки WeasyPrint для лучшего качества
            weasy_options = {}
            
            # Добавляем кастомные шрифты если указаны
            if custom_fonts:
                font_config = self._create_font_config(custom_fonts)
                if font_config:
                    weasy_options['font_config'] = font_config
            
            pdf_bytes = html_doc.write_pdf(**weasy_options)
            
            duration_ms = (time.time() - start_time) * 1000
            file_size = len(pdf_bytes)
            
            LoggingConfig.log_export_operation(
                operation="pdf_export_weasyprint",
                character_id=character.id,
                format_type=format_type,
                export_type="pdf",
                user_id=user_id,
                duration_ms=duration_ms,
                file_size=file_size,
                success=True
            )
            
            return pdf_bytes
            
        except Exception as e:
            # Fallback на ReportLab при ошибке WeasyPrint
            LoggingConfig.get_export_logger().warning(
                f"WeasyPrint failed, falling back to ReportLab: {str(e)}"
            )
            return await self._export_pdf_with_reportlab(
                character, checklists, format_type, user_id, start_time
            )
    
    def _load_theme_css(self, theme: str) -> str:
        """Загружает CSS для указанной темы."""
        if theme == "default":
            return ""
        
        theme_path = self.template_dir / "themes" / f"{theme}.css"
        if theme_path.exists():
            try:
                return theme_path.read_text(encoding='utf-8')
            except Exception as e:
                LoggingConfig.get_export_logger().warning(
                    f"Failed to load theme {theme}: {str(e)}"
                )
        
        return ""
    
    def _create_font_config(self, custom_fonts: Dict[str, str]):
        """Создает конфигурацию шрифтов для WeasyPrint."""
        try:
            # Пока отключаем кастомные шрифты из-за проблем совместимости
            LoggingConfig.get_export_logger().info(
                "Custom fonts temporarily disabled for compatibility"
            )
            return None
        except Exception as e:
            LoggingConfig.get_export_logger().warning(
                f"Failed to create font config: {str(e)}"
            )
            return None
    
    async def _export_pdf_with_reportlab(
        self,
        character: Character,
        checklists: list[ChecklistWithResponses],
        format_type: str,
        user_id: Optional[int],
        start_time: float
    ) -> bytes:
        """Экспорт PDF с использованием ReportLab (оригинальный метод)."""
        buffer = io.BytesIO()
        
        # Создание PDF документа
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Стили
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Heading1'],
            fontName=self.default_font,
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Центр
        ))
        
        styles.add(ParagraphStyle(
            name='CustomHeading',
            parent=styles['Heading2'],
            fontName=self.default_font,
            fontSize=16,
            spaceAfter=12
        ))
        
        styles.add(ParagraphStyle(
            name='CustomNormal',
            parent=styles['Normal'],
            fontName=self.default_font,
            fontSize=12,
            spaceAfter=6
        ))
        
        styles.add(ParagraphStyle(
            name='CustomBold',
            parent=styles['Normal'],
            fontName=self.default_font,
            fontSize=12,
            spaceAfter=4,
            textColor=colors.black
        ))
        
        # Контент документа
        story = []
        
        # Заголовок
        title = f"Анализ персонажа: {character.name}"
        story.append(Paragraph(title, styles['CustomTitle']))
        story.append(Spacer(1, 12))
        
        # Базовая информация о персонаже
        story.append(Paragraph("Базовая информация", styles['CustomHeading']))
        
        char_info = [
            ["Имя:", character.name],
            ["Важность:", f"{character.importance_score:.2f}" if character.importance_score else "Не определена"],
            ["Дата анализа:", datetime.now().strftime("%d.%m.%Y %H:%M")],
        ]
        
        if character.aliases:
            char_info.append(["Псевдонимы:", ", ".join(character.aliases)])
        
        # Таблица с информацией
        char_table = Table(char_info, colWidths=[2*inch, 4*inch])
        char_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.grey),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), self.default_font),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('BACKGROUND', (1, 0), (1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(char_table)
        story.append(Spacer(1, 20))
        
        # Определяем пол персонажа (можно расширить логику)
        character_gender = self._detect_character_gender(character)
        
        # Чеклисты
        if format_type == "detailed":
            story.extend(self._add_detailed_checklists_to_pdf(checklists, styles, character_gender))
        elif format_type == "summary":
            story.extend(self._add_summary_checklists_to_pdf(checklists, styles, character_gender))
        else:  # compact
            story.extend(self._add_compact_checklists_to_pdf(checklists, styles, character_gender))
        
        # Генерация PDF
        doc.build(story)
        buffer.seek(0)
        
        duration_ms = (time.time() - start_time) * 1000
        file_size = len(buffer.getvalue())
        
        LoggingConfig.log_export_operation(
            operation="pdf_export_reportlab",
            character_id=character.id,
            format_type=format_type,
            export_type="pdf",
            user_id=user_id,
            duration_ms=duration_ms,
            file_size=file_size,
            success=True
        )
        
        return buffer.getvalue()
    
    def _add_detailed_checklists_to_pdf(self, checklists: list, styles, character_gender: Optional[str] = None) -> list:
        """Добавить детальные чеклисты в PDF."""
        story = []
        
        for checklist in checklists:
            story.append(Paragraph(f"Чеклист: {checklist.title}", styles['CustomHeading']))
            
            if checklist.description:
                story.append(Paragraph(checklist.description, styles['CustomNormal']))
                story.append(Spacer(1, 6))
            
            # Проходим по новой структуре чеклиста: Section → Subsection → QuestionGroup → Question
            for section in checklist.sections:
                if not section.subsections:
                    continue
                    
                story.append(Paragraph(f"📝 {section.title}", styles['CustomHeading']))
                
                for subsection in section.subsections:
                    if subsection.question_groups:
                        for group in subsection.question_groups:
                            if group.title and group.title != subsection.title:
                                story.append(Paragraph(f"🔹 {group.title}", styles['CustomBold']))
                            
                            for question in group.questions:
                                if question.current_response:
                                    story.append(Paragraph(f"• {question.text}", styles['CustomNormal']))
                                    
                                    # Получаем ответ в зависимости от типа и пола персонажа
                                    answer_text = self._get_answer_text(question.current_response, character_gender)
                                    if answer_text:
                                        story.append(Paragraph(answer_text, styles['CustomNormal']))
                                    
                                    # Добавляем совет если есть
                                    if question.current_response.answer and question.current_response.answer.hint:
                                        story.append(Paragraph(f"Совет: {question.current_response.answer.hint}", styles['CustomNormal']))
                                    
                                    # Добавляем упражнения если есть
                                    if question.current_response.answer and question.current_response.answer.exercise:
                                        story.append(Paragraph(f"Упражнение: {question.current_response.answer.exercise}", styles['CustomNormal']))
                                    
                                    # Добавляем комментарий если есть
                                    if question.current_response.comment:
                                        story.append(Paragraph(f"Комментарий: {question.current_response.comment}", styles['CustomNormal']))
                                    
                                    story.append(Spacer(1, 4))
                
                story.append(Spacer(1, 8))
            
            story.append(Spacer(1, 20))
        
        return story
    
    def _add_summary_checklists_to_pdf(self, checklists: list, styles, character_gender: Optional[str] = None) -> list:
        """Добавить краткие чеклисты в PDF."""
        story = []
        
        story.append(Paragraph("Краткий обзор чеклистов", styles['CustomHeading']))
        
        for checklist in checklists:
            # Подсчитываем общее количество вопросов и ответов по новой структуре
            total_questions = 0
            answered_questions = 0
            
            for section in checklist.sections:
                for subsection in section.subsections:
                    for group in subsection.question_groups:
                        for question in group.questions:
                            total_questions += 1
                            if question.current_response and self._has_answer(question.current_response):
                                answered_questions += 1
            
            completion_rate = (answered_questions / total_questions * 100) if total_questions > 0 else 0
            
            summary_data = [
                [checklist.title, f"{answered_questions}/{total_questions}", f"{completion_rate:.1f}%"]
            ]
            
            summary_table = Table(summary_data, colWidths=[3*inch, 1*inch, 1*inch])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, -1), self.default_font),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(summary_table)
            story.append(Spacer(1, 6))
        
        return story
    
    def _add_compact_checklists_to_pdf(self, checklists: list, styles, character_gender: Optional[str] = None) -> list:
        """Добавить компактные чеклисты в PDF."""
        story = []
        
        story.append(Paragraph("Компактный отчет", styles['CustomHeading']))
        
        total_answered = 0
        total_questions = 0
        
        for checklist in checklists:
            # Подсчитываем вопросы из новой структуры checklist
            questions_count = 0
            answered_count = 0
            for section in checklist.sections:
                for subsection in section.subsections:
                    for group in subsection.question_groups:
                        for question in group.questions:
                            questions_count += 1
                            if question.current_response and self._has_answer(question.current_response):
                                answered_count += 1
            total_questions += questions_count
            total_answered += answered_count
        
        overall_completion = (total_answered / total_questions * 100) if total_questions > 0 else 0
        
        story.append(Paragraph(f"Общая статистика: {total_answered}/{total_questions} вопросов ({overall_completion:.1f}%)", styles['CustomNormal']))
        story.append(Spacer(1, 12))
        
        return story
    
    async def export_character_docx(
        self, 
        character: Character, 
        checklists: list[ChecklistWithResponses],
        format_type: str = "detailed",
        user_id: Optional[int] = None
    ) -> bytes:
        """
        Экспорт персонажа в DOCX формат.
        
        Args:
            character: Модель персонажа
            checklists: Список чеклистов с ответами
            format_type: Тип отчета ('detailed', 'summary', 'compact')
            user_id: ID пользователя для логирования
            
        Returns:
            DOCX файл в виде байтов
        """
        start_time = time.time()
        
        try:
            # Создание документа
            doc = Document()
            
            # Заголовок
            title = doc.add_heading(f'Анализ персонажа: {character.name}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Базовая информация
            doc.add_heading('Базовая информация', level=1)
            
            info_table = doc.add_table(rows=3, cols=2)
            info_table.style = 'Table Grid'
            
            info_table.cell(0, 0).text = 'Имя:'
            info_table.cell(0, 1).text = character.name
            
            info_table.cell(1, 0).text = 'Важность:'
            info_table.cell(1, 1).text = f"{character.importance_score:.2f}" if character.importance_score else "Не определена"
            
            info_table.cell(2, 0).text = 'Дата анализа:'
            info_table.cell(2, 1).text = datetime.now().strftime("%d.%m.%Y %H:%M")
            
            if character.aliases:
                row = info_table.add_row()
                row.cells[0].text = 'Псевдонимы:'
                row.cells[1].text = ', '.join(character.aliases)
            
            # Определяем пол персонажа
            character_gender = self._detect_character_gender(character)
            
            # Добавляем чеклисты в зависимости от формата
            if format_type == "detailed":
                self._add_detailed_checklists_to_docx(doc, checklists, character_gender)
            elif format_type == "summary":
                self._add_summary_checklists_to_docx(doc, checklists, character_gender)
            else:  # compact
                self._add_compact_checklists_to_docx(doc, checklists, character_gender)
            
            # Сохранение в байты
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            duration_ms = (time.time() - start_time) * 1000
            file_size = len(buffer.getvalue())
            
            LoggingConfig.log_export_operation(
                operation="docx_export",
                character_id=character.id,
                format_type=format_type,
                export_type="docx",
                user_id=user_id,
                duration_ms=duration_ms,
                file_size=file_size,
                success=True
            )
            
            return buffer.getvalue()
            
        except Exception as e:
            duration_ms = (time.time() - start_time) * 1000
            LoggingConfig.log_export_operation(
                operation="docx_export",
                character_id=character.id,
                format_type=format_type,
                export_type="docx",
                user_id=user_id,
                duration_ms=duration_ms,
                success=False,
                error=str(e)
            )
            
            if "docx" in str(e).lower() or "document" in str(e).lower():
                raise ExportError(
                    "Ошибка создания DOCX документа",
                    details=f"python-docx ошибка: {str(e)}"
                )
            else:
                raise ExportError(
                    "Неизвестная ошибка при создании DOCX",
                    details=str(e)
                )
    
    def _add_detailed_checklists_to_docx(self, doc: Document, checklists: list, character_gender: Optional[str] = None):
        """Добавить детальные чеклисты в DOCX."""
        for checklist in checklists:
            doc.add_heading(f'Чеклист: {checklist.title}', level=1)
            
            if checklist.description:
                doc.add_paragraph(checklist.description)
            
            # Проходим по новой структуре чеклиста: Section → Subsection → QuestionGroup → Question
            for section in checklist.sections:
                if not section.subsections:
                    continue
                    
                doc.add_heading(section.title, level=2)
                
                for subsection in section.subsections:
                    if subsection.question_groups:
                        for group in subsection.question_groups:
                            if group.title and group.title != subsection.title:
                                doc.add_heading(group.title, level=3)
                            
                            for question in group.questions:
                                if question.current_response:
                                    # Вопрос
                                    question_p = doc.add_paragraph()
                                    question_p.add_run('Вопрос: ').bold = True
                                    question_p.add_run(question.text)
                                    
                                    # Ответ
                                    answer_text = self._get_answer_text(question.current_response, character_gender)
                                    if answer_text:
                                        answer_p = doc.add_paragraph()
                                        answer_p.add_run('Ответ: ').bold = True
                                        answer_p.add_run(answer_text)
                                    
                                    # Совет
                                    if question.current_response.answer and question.current_response.answer.hint:
                                        hint_p = doc.add_paragraph()
                                        hint_p.add_run('Совет: ').bold = True
                                        hint_p.add_run(question.current_response.answer.hint)
                                    
                                    # Упражнения
                                    if question.current_response.answer and question.current_response.answer.exercise:
                                        exercise_p = doc.add_paragraph()
                                        exercise_p.add_run('Упражнение: ').bold = True
                                        exercise_p.add_run(question.current_response.answer.exercise)
                                    
                                    # Комментарий
                                    if question.current_response.comment:
                                        comment_p = doc.add_paragraph()
                                        comment_p.add_run('Комментарий: ').bold = True
                                        comment_p.add_run(question.current_response.comment)
                                    
                                    # Источник
                                    if hasattr(question.current_response, 'source_type') and question.current_response.source_type:
                                        source_text = {
                                            "FOUND_IN_TEXT": "Найдено в тексте",
                                            "LOGICALLY_DERIVED": "Логически выведено",
                                            "IMAGINED": "Придумано"
                                        }.get(question.current_response.source_type.value if hasattr(question.current_response.source_type, 'value') else str(question.current_response.source_type), "Неизвестно")
                                        
                                        source_p = doc.add_paragraph()
                                        source_p.add_run('Источник: ').bold = True
                                        source_p.add_run(source_text)
                                    
                                    doc.add_paragraph("")  # Пустая строка для разделения
    
    def _add_summary_checklists_to_docx(self, doc: Document, checklists: list, character_gender: Optional[str] = None):
        """Добавить краткие чеклисты в DOCX."""
        doc.add_heading('Краткий обзор чеклистов', level=1)
        
        summary_table = doc.add_table(rows=1, cols=3)
        summary_table.style = 'Table Grid'
        
        # Заголовки
        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = 'Чеклист'
        hdr_cells[1].text = 'Отвечено'
        hdr_cells[2].text = 'Процент'
        
        for checklist in checklists:
            row_cells = summary_table.add_row().cells
            # Подсчитываем вопросы из новой структуры checklist для DOCX summary
            total_questions = 0
            answered_questions = 0
            for section in checklist.sections:
                for subsection in section.subsections:
                    for group in subsection.question_groups:
                        for question in group.questions:
                            total_questions += 1
                            if question.current_response and self._has_answer(question.current_response):
                                answered_questions += 1
            completion_rate = (answered_questions / total_questions * 100) if total_questions > 0 else 0
            
            row_cells[0].text = checklist.title
            row_cells[1].text = f"{answered_questions}/{total_questions}"
            row_cells[2].text = f"{completion_rate:.1f}%"
    
    def _add_compact_checklists_to_docx(self, doc: Document, checklists: list, character_gender: Optional[str] = None):
        """Добавить компактные чеклисты в DOCX."""
        doc.add_heading('Компактный отчет', level=1)
        
        total_answered = 0
        total_questions = 0
        
        for checklist in checklists:
            # Подсчитываем вопросы из новой структуры checklist
            questions_count = 0
            answered_count = 0
            for section in checklist.sections:
                for subsection in section.subsections:
                    for group in subsection.question_groups:
                        for question in group.questions:
                            questions_count += 1
                            if question.current_response and self._has_answer(question.current_response):
                                answered_count += 1
            total_questions += questions_count
            total_answered += answered_count
        
        overall_completion = (total_answered / total_questions * 100) if total_questions > 0 else 0
        
        doc.add_paragraph(f"Общая статистика: {total_answered}/{total_questions} вопросов ({overall_completion:.1f}%)")


    def _get_answer_text(self, response, character_gender: Optional[str] = None) -> Optional[str]:
        """
        Получить текст ответа из ChecklistResponse с учетом новой архитектуры.
        
        Args:
            response: ChecklistResponse объект
            character_gender: Пол персонажа ('male', 'female') для выбора правильного значения
            
        Returns:
            Текст ответа или None
        """
        # Если есть связанный ответ (ChecklistAnswer)
        if hasattr(response, 'answer') and response.answer:
            # Выбираем значение в зависимости от пола персонажа
            if character_gender == 'female':
                # Приоритет: exported_value_female -> value_female -> exported_value_male -> value_male
                if response.answer.exported_value_female:
                    return response.answer.exported_value_female
                elif response.answer.value_female:
                    return response.answer.value_female
                elif response.answer.exported_value_male:
                    return response.answer.exported_value_male
                elif response.answer.value_male:
                    return response.answer.value_male
            else:
                # Приоритет: exported_value_male -> value_male -> exported_value_female -> value_female
                if response.answer.exported_value_male:
                    return response.answer.exported_value_male
                elif response.answer.value_male:
                    return response.answer.value_male
                elif response.answer.exported_value_female:
                    return response.answer.exported_value_female
                elif response.answer.value_female:
                    return response.answer.value_female
        
        # Если есть текстовый ответ (свободный ввод)
        if hasattr(response, 'answer_text') and response.answer_text:
            return response.answer_text
        
        # Fallback на старое поле (для совместимости)
        if hasattr(response, 'answer') and isinstance(response.answer, str):
            return response.answer
        
        return None
    
    def _has_answer(self, response) -> bool:
        """
        Проверить, есть ли ответ в ChecklistResponse.
        
        Args:
            response: ChecklistResponse объект
            
        Returns:
            True если есть ответ, False иначе
        """
        return self._get_answer_text(response) is not None
    
    def _detect_character_gender(self, character: Character) -> Optional[str]:
        """
        Определить пол персонажа на основе доступных данных.
        
        Args:
            character: Объект персонажа
            
        Returns:
            'male', 'female' или None если не удалось определить
        """
        # Используем поле gender из модели персонажа
        if hasattr(character, 'gender') and character.gender:
            # Проверяем разные способы получения значения enum
            gender_value = None
            if hasattr(character.gender, 'value'):
                gender_value = character.gender.value
            else:
                gender_value = str(character.gender).lower()
            
            if gender_value == 'male':
                return 'male'
            elif gender_value == 'female':
                return 'female'
            elif 'male' in gender_value:
                return 'male'
            elif 'female' in gender_value:
                return 'female'
        
        # Fallback: анализ имени персонажа (простая эвристика)
        if character.name:
            name_lower = character.name.lower()
            # Простые женские окончания в русских именах
            female_endings = ['а', 'я', 'ия', 'ья', 'на', 'ка']
            if any(name_lower.endswith(ending) for ending in female_endings):
                return 'female'
        
        # По умолчанию возвращаем мужской пол
        return 'male'


# Создание экземпляра сервиса
export_service = ExportService()
